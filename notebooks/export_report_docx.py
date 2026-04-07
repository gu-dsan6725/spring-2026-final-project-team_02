"""Export progress_report.md to a Google-Docs-ready .docx file."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, copy

BASE   = os.path.dirname(__file__) + "/.."
PLOTS  = BASE + "/results/plots"
OUT    = BASE + "/docs/progress_report.docx"

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge, color in kwargs.items():
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"),   "single")
        e.set(qn("w:sz"),    "4")
        e.set(qn("w:color"), color)
        tcBorders.append(e)
    tcPr.append(tcBorders)

def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
from docx.shared import Inches as _I
for section in doc.sections:
    section.top_margin    = _I(0.9)
    section.bottom_margin = _I(0.9)
    section.left_margin   = _I(1.0)
    section.right_margin  = _I(1.0)

# ── Default font ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── Helpers ───────────────────────────────────────────────────────────────────
DARK   = RGBColor(0x2C, 0x3E, 0x50)
ACCENT = RGBColor(0x4A, 0x90, 0xD9)
MUTED  = RGBColor(0x55, 0x55, 0x55)
GREEN  = RGBColor(0x27, 0xAE, 0x60)
RED    = RGBColor(0xE7, 0x4C, 0x3C)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = DARK
    p.runs[0].font.size = Pt(16)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = ACCENT
    p.runs[0].font.size = Pt(13)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)

def body(text, bold=False, italic=False, color=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = "Calibri"
    run.font.size  = Pt(11)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)

def insert_image(path, width=6.0):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = MUTED
    p.paragraph_format.space_after = Pt(10)

def make_table(headers, rows, header_hex="2C3E50", alt_hex="EAF2FB"):
    cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=cols)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        c.paragraphs[0].runs[0].font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(c, header_hex)

    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        bg = alt_hex if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            c.text = str(val)
            c.paragraphs[0].runs[0].font.size = Pt(10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(c, bg)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t

def divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("HERALD — Milestone 2 Progress Report")
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = DARK

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Hierarchical Escalating Retrieval and LLM Disagreement Pipeline")
r2.italic = True
r2.font.size = Pt(12)
r2.font.color.rgb = MUTED

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("Data Preparation & Initial Agent Implementation  |  2026-03-15")
r3.font.size = Pt(11)
r3.font.color.rgb = MUTED

divider()

# ══════════════════════════════════════════════════════════════════════════════
# 1. WHAT HERALD DOES
# ══════════════════════════════════════════════════════════════════════════════
h1("1. What HERALD Does")
body(
    "HERALD validates outputs from economics research LLM agents before they enter reports or "
    "downstream reasoning. Claims are routed through progressively expensive validators — "
    "stopping as soon as any tier is confident — balancing cost, latency, and accuracy."
)
body("Five checkpoint types are validated:", space_after=4)

make_table(
    ["CP", "Type", "What is checked"],
    [
        ["1", "Retrieval",          "Retrieved document is topically relevant to the query"],
        ["2", "Claim Extraction",   "Extracted claim directly entailed by source; numbers exact"],
        ["3", "Synthesis",          "Paragraph faithfully represents all contributing claims"],
        ["4", "Numerical",          "Numbers within accepted rounding; direction correct"],
        ["5", "Causal",             "Causal language does not overstate source's evidence strength"],
    ]
)
body("Verdicts: VALID / INVALID / UNCERTAIN (→ human review).", italic=True, color=MUTED, space_after=2)

# ══════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("2. System Architecture")
insert_image(f"{PLOTS}/architecture.png", width=6.2)
caption("Figure 1 — HERALD four-tier escalation pipeline. Each tier runs only if the previous tier's confidence falls below its threshold.")
body(
    "Every case returns a complete EscalationPacket with verdicts, confidence scores, and reasoning "
    "from every tier that ran — providing a full audit trail. Thresholds T1 = 0.70 and T2 = 0.80 "
    "are configurable in configs/default.yaml without code changes."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. DATA PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Data Pipeline")

h2("3.1 Datasets")
make_table(
    ["Dataset", "Cases", "Purpose"],
    [
        ["sample_cases.json",          "10",  "Labeled eval set (ground truth)"],
        ["feasibility_samples.json",   "40",  "Expanded NLI feasibility / smoke-test set"],
        ["unlabeled_pairs.json",       "30",  "Source/claim pairs for Groq labeling"],
        ["weak_labeled.json",          "~40", "Groq-generated weak labels for fine-tuning"],
    ]
)
body(
    "All cases are grounded in U.S. macroeconomics (Fed policy, inflation, housing, labor markets) "
    "with three-way labels: valid / invalid / ambiguous. Validity criteria for each checkpoint type "
    "are precisely defined — for example, CP4 (Numerical) allows ±1 unit of rounding ('1.08M' from "
    "source '1.078M' is valid; '5% decline' from '3.2%' is not)."
)

h2("3.2 Feasibility Gate — NLI Signal Verification")
body(
    "Before building the full pipeline, a GO/NO-GO gate verified that the base DeBERTa model "
    "produces meaningful discriminative signal on economics data without any fine-tuning."
)
insert_image(f"{PLOTS}/nli_separation.png", width=5.4)
caption("Figure 2 — DeBERTa entailment/contradiction scores by label (40 cases). The 0.545 entailment gap confirms Tier 1 can discriminate without fine-tuning.")
body(
    "Gate passed. Fine-tuning (script notebooks/finetune_tier1.py is ready) is repositioned as "
    "a Phase 2 improvement once 500+ weak-labeled pairs exist — not a Phase 1 prerequisite. "
    "This unblocked the pipeline 2–3 weeks earlier than planned.",
    italic=True
)

# ══════════════════════════════════════════════════════════════════════════════
# 4. AGENT IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Agent Implementation")

h2("Tier 1 — NLI Classifier")
body(
    "Runs locally on CPU; zero cost. Source context is the NLI premise, agent output is the "
    "hypothesis. DeBERTa produces entailment / contradiction / neutral probabilities; the "
    "highest-scoring class becomes the verdict subject to the T1 threshold (0.70). For multi-source "
    "checkpoints (CP1, CP3), three aggregation strategies are supported: max_entailment, "
    "max_contradiction, and mean."
)

h2("Tier 2 — LLM Judge")
body(
    "Single Groq call to llama-3.3-70b-versatile (free tier, temp=0.1). The judge receives the "
    "output text, source context, and Tier 1's raw NLI scores. It returns structured JSON "
    "{ verdict, confidence, reasoning, key_issues } enforced via Groq's json_object response format. "
    "If stated confidence is below T2 (0.80), the verdict is overridden to UNCERTAIN and the case "
    "escalates. Three-attempt retry with 2s backoff handles rate limits."
)

h2("Tier 3 — Multi-Agent Debate")
body(
    "Three sequential Groq calls (~6s total). The Advocate builds the strongest case the output "
    "is valid (temp=0.3); the Critic argues it is not (temp=0.3); the Judge rules on source "
    "evidence — not rhetorical quality — and returns a structured JSON verdict (temp=0.1). "
    "Both advocate and critic receive the full prior analysis (Tier 1 scores + Tier 2 reasoning), "
    "targeting the debate at exactly the point of disagreement."
)
body(
    "Design note: the 0.3 / 0.1 temperature split was added after testing showed near-zero "
    "temperature produces near-identical advocate/critic arguments, collapsing the debate.",
    italic=True, color=MUTED
)

h2("Tier 4 — Human Review")
body(
    "When all automated tiers remain uncertain, the pipeline generates a structured JSON review "
    "packet: full reasoning trace from all prior tiers, a framed question tailored to the source "
    "of uncertainty, and all context needed for review. Packet generation is complete; "
    "a web portal for verdict submission is the primary remaining gap (see Risks)."
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. RESULTS
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Results")

h2("5.1 Summary — 10 Labeled Cases")
make_table(
    ["Metric", "Value"],
    [
        ["Overall accuracy",              "90% (9/10)"],
        ["Resolved at Tier 1 (NLI only)", "7 cases — 85.7% accuracy"],
        ["Resolved at Tier 2 (LLM judge)","2 cases — 100% accuracy"],
        ["Resolved at Tier 3 (debate)",   "1 case  — 100% accuracy"],
        ["Human review required",         "0 cases"],
        ["Total cost",                    "$0.00"],
    ]
)

h2("5.2 Escalation Profile and Error Analysis — 40-Case Run")
insert_image(f"{PLOTS}/plot2_escalation_profile.png", width=6.0)
caption("Figure 3 — Escalation profile by checkpoint type. Numerical resolves 71% at Tier 1; synthesis and causal consistently require Tier 2 or Tier 3.")

insert_image(f"{PLOTS}/plot4_confusion_analysis.png", width=6.2)
caption("Figure 4 — Confusion matrix and per-checkpoint error rate. The single misclassification is the ambiguous test case, mis-resolved at Tier 1 with 0.832 confidence.")

body(
    "The confusion matrix shows HERALD never outputs UNCERTAIN on this test set, meaning the "
    "ambiguous case receives a confident wrong verdict at Tier 1 rather than escalating. "
    "This is the expected failure mode and drives the risk mitigations below."
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. ARCHITECTURE UPDATES FROM EARLY LEARNINGS
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Architecture Updates from Early Learnings")

bullet(None, bold_prefix="Base DeBERTa is strong enough to defer fine-tuning. ")
doc.paragraphs[-1].runs[-1].text += (
    "The feasibility gate showed a 0.545 entailment gap without domain adaptation. "
    "Fine-tuning is now a Phase 2 improvement, unblocking the pipeline 2–3 weeks early."
)

bullet(None, bold_prefix="Synthesis and causal checkpoints need a dedicated high-confidence path. ")
doc.paragraphs[-1].runs[-1].text += (
    "The escalation profile confirms 50–57% of CP3/CP5 cases route through Tier 2 or Tier 3 "
    "(vs. 29% for numerical). Tier 3 debate is now designed as the expected resolution path "
    "for these types, not a rare fallback."
)

bullet(None, bold_prefix="Debate quality requires temperature differentiation. ")
doc.paragraphs[-1].runs[-1].text += (
    "Advocate and critic run at temp=0.3 to produce meaningfully opposed arguments; "
    "the judge uses temp=0.1 for consistent verdicts."
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. RISKS AND MITIGATION PLANS
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Risks and Mitigation Plans")

risks = [
    (
        "Groq rate limits interrupt multi-case runs", "High",
        "Two of four deliverable plots failed mid-run due to 429 errors. "
        "Retry logic (3 attempts, 2s backoff) is in place. Additional: result caching by checkpoint hash, "
        "configurable inter-request delay, Groq paid tier as fallback (~$5 for full 40-case sweep)."
    ),
    (
        "DeBERTa confidence miscalibrated for ambiguous cases", "High",
        "The ambiguous test case resolved at Tier 1 with 0.832 confidence instead of escalating. "
        "Mitigation: run calibration_analysis.py to compute ECE; consider checkpoint-type-specific "
        "thresholds (lower T1 for CP3/CP5) to force borderline cases through higher tiers."
    ),
    (
        "No Tier 4 review portal", "Medium",
        "The 0% human review rate reflects a mostly clear-cut test set. Synthesis/causal cases "
        "in production will require review. Packet generation is done; a minimal Streamlit or "
        "Flask form is the next build priority."
    ),
    (
        "Weak label quality contaminates fine-tuning", "Medium",
        "LLM-generated labels may be systematically wrong for subtle causal overstatement. "
        "Mitigation: filter to confidence ≥ 0.85 before training; A/B evaluate fine-tuned "
        "vs. base model on the hand-labeled 10-case set before deploying."
    ),
]

make_table(
    ["Risk", "Severity", "Mitigation"],
    [[r[0], r[1], r[2]] for r in risks],
    header_hex="2C3E50",
    alt_hex="FDFEFE"
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. STATUS
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Status Summary")

make_table(
    ["Component", "Status"],
    [
        ["Full 4-tier escalation pipeline",        "Complete"],
        ["Tier 1 NLI classifier (DeBERTa, local)", "Complete"],
        ["Tier 2 LLM judge (Groq)",                "Complete"],
        ["Tier 3 multi-agent debate",              "Complete"],
        ["Tier 4 human review packet generation",  "Complete"],
        ["Labeled datasets (50 cases total)",      "Complete"],
        ["Weak label generation",                  "Complete"],
        ["Evaluation framework",                   "Complete"],
        ["Threshold sweep / baseline comparison",  "Complete"],
        ["DeBERTa fine-tuning",                    "Pending (script ready; needs 500+ pairs)"],
        ["Tier 4 review portal UI",                "Not started"],
        ["2 remaining deliverable plots",          "Blocked by Groq rate limits"],
    ],
    header_hex="2C3E50",
    alt_hex="EAF2FB"
)

divider()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Branch: manav-dev  |  Commit: b35bf2d  |  Herald v1.1")
r.font.size = Pt(9)
r.font.color.rgb = MUTED
r.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
